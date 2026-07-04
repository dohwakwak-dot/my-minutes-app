import streamlit as st
from google import genai
from google.genai import types
from docx import Document
import io
import re

# 웹 화면 레이아웃 설정
st.set_page_config(page_title="엔지니어링 스마트 회의록 시스템", layout="wide")

# 🔒 화면 왼쪽(사이드바)에서 각자 본인의 API 키를 입력하도록 설정하여 사용자 토큰을 방어합니다.
with st.sidebar:
    st.header("🔑 API Key 설정")
    GEMINI_API_KEY = st.text_input("Gemini API Key 입력", type="password", placeholder="AIzaSy...")
    
    st.markdown("""
    ### 🔰 초간단 API Key 발급 방법
    1. [👉 구글 AI 스튜디오 바로가기](https://aistudio.google.com/) 클릭
    2. 소지하신 **구글 계정(Gmail)**으로 로그인
    3. 약관 창이 뜨면 모두 체크 후 **[Accept/Continue]** 클릭
    4. 좌측 상단의 파란색 **`[Get API key]`** 버튼 클릭
    5. **`[Create API key]`** ➡️ **`[Create API key in new project]`** 클릭
    6. 생성된 `AIzaSy...`로 시작하는 긴 키를 복사(`Copy`)해서 위 칸에 붙여넣기 하시면 됩니다!
    """)
    st.write("---")
    st.caption("※ 本 프로그램은 입력된 Key를 수집하지 않으며, 해당 계정의 무료/유료 토큰만 소모합니다.")

# 입력한 키로 구글 재미나이 가동하기
if GEMINI_API_KEY:
    client = genai.Client(api_key=GEMINI_API_KEY)
else:
    client = None

# 웹 화면 메인 타이틀 꾸미기
st.title("🏗️ 엔지니어링 맞춤형 스마트 회의록 자동화 시스템")
st.write("발언자 그룹 매칭 단계를 거쳐 100% 정확한 사실 기반 회의록을 도출합니다.")

# 1단계: 참석자 및 분야별 정보 입력받기
st.subheader("👥 1단계: 회의 참석자 정보 입력")
user_count = st.number_input("정보를 입력할 참석 인원수 (명)", min_value=1, value=3, step=1)

member_data = []
for i in range(int(user_count)):
    col_name, col_field = st.columns(2)
    with col_name:
        name = st.text_input(f"👤 참석자 {i+1} 이름", key=f"name_{i}", placeholder="홍길동")
    with col_field:
        field = st.text_input(f"⚙️ 참석자 {i+1} 담당 분야/직급", key=f"field_{i}", placeholder="예: 상하수도, 구조부, 발주청 감독관 등")
    if name:
        member_data.append({"이름": name, "분야": field if field else "미지정"})

# 선택 박스용 실제 이름 목록 리스트 생성
member_names_list = [m['이름'] for m in member_data]

st.write("---")

# 2단계: 회의 메모 사전 입력 (덤핑식)
st.subheader("📝 2단계: 회의 중 메모한 내용 입력")
memo_input = st.text_area("회의 메모 덤핑 입력란", height=120, placeholder="예: 상하수도 관경 변경 검토 요청함. 발주청 지시사항으로 다음 주까지 공정표 제출 바람...")

st.write("---")

# 3단계: 녹음 파일 업로드 및 가생성
st.subheader("🎵 3단계: 회의 녹음 파일 업로드")
uploaded_file = st.file_uploader("회의 파일(mp3, wav 등)을 올려주세요.", type=["mp3", "wav", "m4a", "wma"])

st.write("---")

# 상태 저장을 위한 세션 설정
if "step1_done" not in st.session_state:
    st.session_state.step1_done = False
if "raw_transcript" not in st.session_state:
    st.session_state.raw_transcript = ""
if "detected_speakers" not in st.session_state:
    st.session_state.detected_speakers = []

# 파일이 올라왔을 때 작동
if uploaded_file is not None and member_data:
    if client is None:
        st.warning("👈 왼쪽 사이드바에 본인의 Gemini API Key를 먼저 입력해 주세요!")
    else:
        # [변경 프로세스] 최종 작성 전, 임시로 화자를 분리하는 버튼
        if not st.session_state.step1_done:
            if st.button("🔍 1차 음성 분석 및 발언자 분리 시작"):
                with st.spinner("재미나이가 음성을 분석하여 목소리 그룹(발언자1, 발언자2...)을 식별하고 있습니다..."):
                    try:
                        audio_bytes = uploaded_file.read()
                        file_ext = uploaded_file.name.split('.')[-1].lower()
                        mime_type = f"audio/{file_ext}" if file_ext in ['mp3', 'wav', 'm4a', 'wma'] else "audio/mp3"
                        audio_part = types.Part.from_bytes(data=audio_bytes, mime_type=mime_type)
                        
                        # AI에게 이름을 유추하지 말고 오직 발언자1, 발언자2 형태로 대화 원본만 뽑아내라고 명시
                        speaker_prompt = """
                        위 오디오 파일의 대화 원본을 토대로 녹취록을 작성해라. 
                        단, 발언자의 실제 이름을 절대 임의로 추측해서 적지 마라. 
                        목소리 톤과 순서에 따라 철저히 '발언자 1', '발언자 2', '발언자 3' 형태로만 구분하여 '발언자 X: 대화내용' 양식으로 대화록 원본 전체를 출력해라.
                        """
                        response = client.models.generate_content(
                            model='gemini-2.5-flash',
                            contents=[audio_part, speaker_prompt]
                        )
                        st.session_state.raw_transcript = response.text
                        
                        # 결과물에서 '발언자 1', '발언자 2' 등의 패턴 추출하기
                        finders = re.findall(r'(발언자\s*\d+)', st.session_state.raw_transcript)
                        st.session_state.detected_speakers = sorted(list(set([f.replace(" ", "") for f in finders])))
                        
                        if not st.session_state.detected_speakers:
                            # 혹시 한글 인식이 튀었을 경우를 대비한 가상 기본 목록 기본 세팅
                            st.session_state.detected_speakers = ["발언자1", "발언자2", "발언자3", "발언자4"]
                            
                        st.session_state.step1_done = True
                        st.rerun()
                    except Exception as e:
                        st.error(f"음성 분석 중 에러가 발생했습니다: {e}")

        # 임시 분석 완료 후 사용자에게 매칭 창을 열어줌
        if st.session_state.step1_done:
            st.subheader("🔄 4단계: 시스템이 감지한 발언자와 실제 참석자 매칭")
            st.success("✅ AI가 음성 파일 속 목소리 그룹을 분리했습니다! 아래 칸에서 각각 매칭되는 실제 사람을 선택해 주세요.")
            
            # 딕셔너리를 사용하여 동적으로 매칭 창 생성
            mapping_results = {}
            
            # 발언자별로 실제 참석자를 고를 수 있는 선택박스 제공
            for speaker in st.session_state.detected_speakers:
                mapping_results[speaker] = st.selectbox(
                    f"🔊 {speaker}는 실제 누구인가요?",
                    options=member_names_list,
                    key=f"map_{speaker}"
                )
            
            st.write("---")
            st.subheader("🚀 5단계: 최종 사실 기반 회의록 최종 생성")
            
            if st.button("✨ 최종 회의록 및 한글 파일 발행"):
                with st.spinner("설정하신 이름 매칭 정보를 바탕으로 최종 완벽한 보고서를 작성 중입니다..."):
                    
                    # 매칭 데이터를 텍스트로 변환하여 AI에게 주입할 준비
                    mapping_str = "\n".join([f"- 변경 전: {k} -> 변경 후 실제 이름: {v}" for k, v in mapping_results.items()])
                    member_info_str = ", ".join([f"{m['이름']}({m['분야']})" for m in member_data])
                    
                    final_prompt = f"""
                    너는 엔지니어링 설계 회사의 베테랑 전문 비서 요원이야.
                    제공된 데이터와 [🚨 작성 규칙]을 100% 준수하여 한글 파일 보고서용 회의록을 완성해줘.

                    [제공된 데이터 자료]
                    1. 1차 분리된 대화 원본:
                    {st.session_state.raw_transcript}
                    
                    2. 작성자가 지정한 정확한 화자 매칭 정보:
                    {mapping_str}
                    
                    3. 전체 인원 및 소속 분야:
                    {member_info_str}
                    
                    4. 작성자가 입력한 회의 사전 메모:
                    {memo_input}

                    [🚨 작성 규칙 - 사실 입각 및 서식 지정]
                    1. 1차 대화 원본에서 '발언자 X'로 표기된 부분들을 제공된 '화자 매칭 정보'를 대조하여 실제 이름으로 완전히 치환해라. 
                    2. 오직 대화 내용과 사전 메모에 명시된 사실로만 요약해라. AI 너의 주관적인 의견, 기술 제안, 예측 아이디어는 '절대' 포함하지 마라.
                    3. 샾(#), 별표(**), 대시(---) 같은 마크다운 기호는 전면 금지한다. 줄바꿈과 공백으로만 가독성을 높여라.
                    4. 문장 끝은 전문 보고서 어조인 명사형 종결어미(-함, -임, - 바람, - 결정됨)를 사용해라.

                    [📋 출력 양식 구조 가이드]
                    [회의 개요]
                    • 참석자 명단: {member_info_str}

                    [주요 안건 및 논의 내용]
                    (안건별로 명확히 단락을 나누고, 세부 논의는 대시(-) 기호의 개조식으로 요약할 것. 문장 끝에는 바뀐 실제 발언자 이름을 괄호 표기할 것.)
                    예시: 안건 1. 유입펌프장 관로 설계 변경 건
                      - 비상시 연계 가능한 관로 설치 반영 완료함 (곽상호)

                    [결론: 분야별 후속 진행 사항]
                    (회의 결과 이후 조치해야 할 태스크를 아래 필수 고려 부서 및 발주청을 검토하여 분리할 것. 담당자가 매칭되어 있다면 담당자 이름도 함께 적어 가이드할 것.)
                    ※ 필수 고려 목록: 상하수도분야, 구조분야, 토질분야, 건축분야, 기계분야, 전기 및 계측제어분야, 인허가 분야, 기타분야, 발주청
                    예시: • [상하수도분야 / 곽상호] 연계 관로 수리계산서 피드백 요청 건 대응 예정

                    [회의 대화 내용 원본 (녹취록)]
                    (1차 분리된 대화 원본의 '발언자 X'를 실제 매칭된 이름으로 전부 수정한 대화록 전체를 그대로 출력할 것. 양식은 반드시 '이름: 대화내용' 순서로 정렬할 것.)
                    """
                    
                    try:
                        # 텍스트 컨텍스트만 활용하여 정밀 가공 요청
                        final_response = client.models.generate_content(
                            model='gemini-2.5-flash',
                            contents=final_prompt
                        )
                        final_minutes = final_response.text
                        st.success("✅ 완벽하게 매칭된 최종 회의록이 발행되었습니다!")
                        
                        # 최종 결과 미리보기창
                        st.subheader("📝 최종 회의록 및 매칭 대화록 미리보기")
                        st.text_area("최종 결과물 창 (복사 가능)", value=final_minutes, height=400)
                        
                        # 파일 빌드 및 다운로드
                        doc = Document()
                        doc.add_heading('📋 스마트 회의록 결과 보고서', level=0)
                        doc.add_paragraph(final_minutes)
                        bio_docx = io.BytesIO()
                        doc.save(bio_docx)
                        
                        bio_hwpx = io.BytesIO()
                        bio_hwpx.write(final_minutes.encode('utf-8'))
                        
                        col_dl1, col_dl2 = st.columns(2)
                        with col_dl1:
                            st.download_button("💾 워드(Docx) 파일로 다운로드", data=bio_docx.getvalue(), file_name="엔지니어링_최종회의록.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        with col_dl2:
                            st.download_button("👑 한글(Hwpx) 파일로 다운로드", data=bio_hwpx.getvalue(), file_name="엔지니어링_최종회의록.hwpx", mime="application/octet-stream")
                            
                        # 새로고침하여 처음부터 다시 할 수 있도록 리셋 버튼 제공
                        if st.button("🔄 처음부터 다시 작성하기"):
                            st.session_state.step1_done = False
                            st.rerun()
                            
                    except Exception as e:
                        st.error(f"최종 조율 중 에러 발생: {e}")
else:
    st.info("💡 왼쪽 사이드바에 API 키를 넣고, 1~3단계 정보를 채우시면 분석 버튼이 나타납니다.")
    st.session_state.step1_done = False
